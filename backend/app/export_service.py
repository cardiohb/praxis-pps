from docx import Document
from io import BytesIO
from app.models import Process

def export_to_word(process: Process) -> bytes:
    doc = Document()
    
    # Header
    doc.add_heading(process.name, level=1)
    doc.add_paragraph(f"Kategorie: {process.category}")
    doc.add_paragraph(f"Verantwortlich: {process.responsible}")
    
    # Schritte
    for i, step in enumerate(process.detailedSteps):
        doc.add_heading(f"Schritt {i+1}: {step.title}", level=2)
        doc.add_paragraph(step.description)
        
        if step.clickPath:
            doc.add_paragraph(f"Klickpfad: {' → '.join(step.clickPath)}")
        
        if step.expectedResult:
            doc.add_paragraph(f"Ergebnis: {step.expectedResult}")
        
        if step.duration:
            doc.add_paragraph(f"Dauer: {step.duration}")
        
        doc.add_paragraph() # Leerer Absatz als Trenner
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
